# -*- coding: utf-8 -*-
"""把两份 markdown 交付物转成排版讲究的 docx（中文宋体正文 + 黑体标题 + 表格/引用/书目悬挂缩进）。
依赖 python-docx。用法：python3 make_docx.py"""
import re, os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
INK = RGBColor(0x22, 0x20, 0x1c)
ACC = RGBColor(0x8a, 0x3b, 0x2e)
ACC2 = RGBColor(0x3d, 0x5a, 0x6c)
GRAY = RGBColor(0x6b, 0x64, 0x5a)

CN_BODY, CN_HEAD, CN_QUOTE = "宋体", "黑体", "楷体"
EN_BODY, EN_MONO = "Times New Roman", "Consolas"


def set_font(run, cn=CN_BODY, en=EN_BODY, size=10.5, bold=False, italic=False, color=INK):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    run.font.name = en
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), cn)


def add_inline(p, text, base_size=10.5, base_cn=CN_BODY, base_color=INK, base_bold=False):
    """解析 **bold** 与 `code`（code 渲染为灰色小字标签）。"""
    tokens = re.split(r'(\*\*.+?\*\*|`[^`]+`)', text)
    for tk in tokens:
        if not tk:
            continue
        if tk.startswith('**') and tk.endswith('**'):
            r = p.add_run(tk[2:-2])
            set_font(r, cn=base_cn, size=base_size, bold=True, color=base_color)
        elif tk.startswith('`') and tk.endswith('`'):
            r = p.add_run(tk[1:-1])
            set_font(r, cn=CN_BODY, en=EN_BODY, size=base_size - 1.5, color=GRAY)
        else:
            r = p.add_run(tk)
            set_font(r, cn=base_cn, size=base_size, bold=base_bold, color=base_color)


def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def build(md_path, out_path):
    doc = Document()
    # 页面 A4 + 边距
    sec = doc.sections[0]
    sec.page_height, sec.page_width = Cm(29.7), Cm(21.0)
    sec.top_margin = sec.bottom_margin = Cm(2.5)
    sec.left_margin = sec.right_margin = Cm(2.6)
    # Normal 默认宋体五号
    nf = doc.styles['Normal'].font
    nf.name = EN_BODY
    nf.size = Pt(10.5)
    doc.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), CN_BODY)

    lines = open(md_path, encoding='utf-8').read().split('\n')
    n = len(lines)
    i = 0
    heading_seen = 0
    while i < n:
        line = lines[i]
        s = line.strip()

        # 跳过空行 / 分隔线
        if not s or s == '---':
            i += 1
            continue

        # 代码块
        if s.startswith('```'):
            i += 1
            buf = []
            while i < n and not lines[i].strip().startswith('```'):
                buf.append(lines[i])
                i += 1
            i += 1
            for cl in buf:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.4)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                p.paragraph_format.space_after = Pt(0)
                r = p.add_run(cl if cl else ' ')
                set_font(r, cn=CN_BODY, en=EN_MONO, size=9, color=GRAY)
            doc.add_paragraph().paragraph_format.space_after = Pt(4)
            continue

        # 表格
        if s.startswith('|') and i + 1 < n and re.match(r'^\|[\s:|\-]+\|', lines[i + 1].strip()):
            rows = []
            while i < n and lines[i].strip().startswith('|'):
                rows.append(lines[i].strip())
                i += 1
            cells = [[c.strip() for c in r.strip('|').split('|')] for r in rows]
            header = cells[0]
            body = cells[2:]
            tbl = doc.add_table(rows=1, cols=len(header))
            tbl.style = 'Table Grid'
            tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for j, h in enumerate(header):
                c = tbl.rows[0].cells[j]
                c.paragraphs[0].text = ''
                add_inline(c.paragraphs[0], h, base_size=10, base_cn=CN_HEAD, base_bold=True)
                shade(c, 'efe9dd')
            for brow in body:
                cellrow = tbl.add_row().cells
                for j, val in enumerate(brow):
                    if j >= len(cellrow):
                        break
                    cellrow[j].paragraphs[0].text = ''
                    add_inline(cellrow[j].paragraphs[0], val, base_size=10)
            doc.add_paragraph().paragraph_format.space_after = Pt(4)
            continue

        # 引用块：每行各成一段
        if s.startswith('>'):
            buf = []
            while i < n and lines[i].strip().startswith('>'):
                buf.append(lines[i].strip().lstrip('>').strip())
                i += 1
            buf = [q for q in buf if q]
            for qi, q in enumerate(buf):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.6)
                p.paragraph_format.space_before = Pt(2 if qi == 0 else 0)
                p.paragraph_format.space_after = Pt(6 if qi == len(buf) - 1 else 1)
                p.paragraph_format.line_spacing = 1.4
                add_inline(p, q, base_size=9.5, base_cn=CN_QUOTE, base_color=GRAY)
            continue

        # 标题
        m = re.match(r'^(#{1,6})\s+(.*)$', s)
        if m:
            level = len(m.group(1))
            text = m.group(2)
            heading_seen += 1
            p = doc.add_paragraph()
            if level == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(6)
                add_inline(p, text, base_size=19, base_cn=CN_HEAD, base_color=ACC, base_bold=True)
            elif level == 2 and heading_seen == 2:   # 副标题
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(10)
                add_inline(p, text, base_size=14, base_cn=CN_HEAD, base_color=ACC2, base_bold=True)
            else:
                sizes = {2: 15, 3: 13, 4: 12, 5: 11.5, 6: 11}
                p.paragraph_format.space_before = Pt(12 if level == 2 else 8)
                p.paragraph_format.space_after = Pt(4)
                add_inline(p, text, base_size=sizes.get(level, 11),
                           base_cn=CN_HEAD, base_color=(ACC if level == 2 else INK), base_bold=True)
            i += 1
            continue

        # 有序列表
        mo = re.match(r'^(\d+)\.\s+(.*)$', s)
        if mo:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.9)
            p.paragraph_format.first_line_indent = Cm(-0.9)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.4
            add_inline(p, f"{mo.group(1)}. {mo.group(2)}")
            i += 1
            continue

        # 无序列表（含书目条目）→ Word 原生项目符号列表（List Bullet 样式）
        if s.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Cm(0.75)
            p.paragraph_format.first_line_indent = Cm(-0.4)
            p.paragraph_format.space_after = Pt(1.5)
            p.paragraph_format.line_spacing = 1.35
            add_inline(p, s[2:])
            i += 1
            continue

        # 缩进续行（解题描述）
        if line.startswith('   ') and s:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.9)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.45
            add_inline(p, s, base_color=GRAY)
            i += 1
            continue

        # 整行加粗 → 小标题（如报刊文章「晚清（18）」、统计行无）
        if re.match(r'^\*\*[^*]+\*\*$', s):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(7)
            p.paragraph_format.space_after = Pt(2)
            add_inline(p, s, base_size=11.5, base_cn=CN_HEAD, base_color=ACC2, base_bold=True)
            i += 1
            continue

        # 普通段落
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.first_line_indent = Cm(0.74)   # 首行缩进两字
        add_inline(p, s, base_size=10.5)
        i += 1

    doc.save(out_path)
    print("written", out_path)


DELIV = os.path.join(BASE, "deliverables")
for src, dst in [("①_资料文献书目.md", "①_资料文献书目.docx"),
                 ("②_资料库技术规划书.md", "②_资料库技术规划书.docx")]:
    build(os.path.join(DELIV, src), os.path.join(DELIV, dst))
